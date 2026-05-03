#!/usr/bin/env python3
"""
Convierte archivos Markdown a formato DOCX (Word).

Uso:
    python scripts/md_to_docx.py <input.md> <output.docx>

Ejemplo:
    python scripts/md_to_docx.py docs/MANUAL_BRIEF_SEMANAL.md ~/Downloads/Manual.docx

Preserva:
    - Encabezados jerarquizados (H1-H6)
    - Tablas con bordes y headers en negrita
    - Listas numeradas y con viñetas
    - Negrita, cursiva, código inline
    - Enlaces clicables (hyperlinks)
    - Bloques de código (Courier New)
    - Citas (blockquotes con indentación)

Requisitos:
    pip install python-docx markdown beautifulsoup4

Convenciones del documento generado:
    - Fuente Arial 11pt
    - Márgenes de 1 pulgada
    - Estilo de tabla: Light Grid Accent 1
"""
import sys
from pathlib import Path

import markdown
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def add_hyperlink(paragraph, url: str, text: str):
    """Agrega un hyperlink al párrafo (no nativo en python-docx)."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',
        is_external=True,
    )
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0563C1')
    rPr.append(color)
    underline = OxmlElement('w:u')
    underline.set(qn('w:val'), 'single')
    rPr.append(underline)
    new_run.append(rPr)
    text_elem = OxmlElement('w:t')
    text_elem.text = text
    new_run.append(text_elem)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_runs(paragraph, element):
    """Recorre el HTML del elemento y agrega runs con formato."""
    for child in element.children:
        if isinstance(child, str):
            if child.strip() or paragraph.runs:
                paragraph.add_run(child)
        elif child.name in ('strong', 'b'):
            run = paragraph.add_run(child.get_text())
            run.bold = True
        elif child.name in ('em', 'i'):
            run = paragraph.add_run(child.get_text())
            run.italic = True
        elif child.name == 'code':
            run = paragraph.add_run(child.get_text())
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
        elif child.name == 'a':
            href = child.get('href', '')
            text = child.get_text()
            if href:
                add_hyperlink(paragraph, href, text)
            else:
                paragraph.add_run(text)
        elif child.name == 'br':
            paragraph.add_run('\n')
        else:
            add_runs(paragraph, child)


def convert(md_path: Path, docx_path: Path):
    """Convierte md_path a docx_path."""
    md_text = md_path.read_text(encoding='utf-8')
    html = markdown.markdown(
        md_text,
        extensions=['extra', 'tables', 'fenced_code', 'sane_lists'],
    )
    soup = BeautifulSoup(html, 'html.parser')

    doc = Document()

    # Estilo base: Arial 11pt
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)

    # Márgenes 1 pulgada
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    for elem in soup.children:
        if not hasattr(elem, 'name') or elem.name is None:
            continue

        if elem.name in ('h1', 'h2', 'h3', 'h4', 'h5', 'h6'):
            level = int(elem.name[1])
            heading = doc.add_heading(elem.get_text(), level=level)
            for run in heading.runs:
                run.font.name = 'Arial'

        elif elem.name == 'p':
            p = doc.add_paragraph()
            add_runs(p, elem)

        elif elem.name == 'ul':
            for li in elem.find_all('li', recursive=False):
                p = doc.add_paragraph(style='List Bullet')
                add_runs(p, li)

        elif elem.name == 'ol':
            for li in elem.find_all('li', recursive=False):
                p = doc.add_paragraph(style='List Number')
                add_runs(p, li)

        elif elem.name == 'table':
            rows_html = elem.find_all('tr')
            if not rows_html:
                continue
            n_cols = max(len(r.find_all(['td', 'th'])) for r in rows_html)
            table = doc.add_table(rows=len(rows_html), cols=n_cols)
            table.style = 'Light Grid Accent 1'
            for ri, tr in enumerate(rows_html):
                cells = tr.find_all(['td', 'th'])
                for ci, td in enumerate(cells):
                    cell = table.cell(ri, ci)
                    cell.text = ''
                    p = cell.paragraphs[0]
                    add_runs(p, td)
                    if td.name == 'th':
                        for run in p.runs:
                            run.bold = True

        elif elem.name == 'pre':
            code = elem.get_text()
            p = doc.add_paragraph()
            run = p.add_run(code)
            run.font.name = 'Courier New'
            run.font.size = Pt(10)

        elif elem.name == 'hr':
            doc.add_paragraph('─' * 60)

        elif elem.name == 'blockquote':
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            add_runs(p, elem)

    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(docx_path)
    print(f'OK: {docx_path}')


if __name__ == '__main__':
    if len(sys.argv) != 3:
        print('Uso: python scripts/md_to_docx.py <input.md> <output.docx>', file=sys.stderr)
        sys.exit(1)
    convert(Path(sys.argv[1]), Path(sys.argv[2]))
