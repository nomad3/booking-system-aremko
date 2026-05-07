"""Genera reporte de follow-ups pendientes y lo envia por email a Jorge.

Una vez por sesion (o ad-hoc) este comando:
1. Carga todas las EncuestaSatisfaccion con requiere_followup=True,
   followup_completado=False (sin filtro de fecha — incluye los atrasados).
2. Para cada caso, llama a Claude Sonnet via OpenRouter para sugerir un
   mensaje WhatsApp empático con la situación específica del cliente.
3. Construye un reporte markdown con todo (data + sugerencia IA).
4. Convierte a .docx (si hay python-docx) o envia el .md directo.
5. Envia por email via Django EmailBackend (SendGrid en produccion).

Uso:
    python manage.py report_pending_followups
    python manage.py report_pending_followups --to otro@email.cl
    python manage.py report_pending_followups --no-email --output /tmp/followups.md
    python manage.py report_pending_followups --no-ai   # skip Claude, solo data
"""
import logging
from datetime import datetime
from io import BytesIO
from pathlib import Path

from django.core.mail import EmailMessage
from django.core.management.base import BaseCommand
from django.utils import timezone

from ventas.models import EncuestaSatisfaccion

logger = logging.getLogger(__name__)


SYSTEM_PROMPT = """Eres un asistente de Aremko Spa Boutique en Puerto Varas, Chile, ayudando a recuperar la relación con clientes que tuvieron una experiencia negativa.

Tu tarea: redactar un mensaje WhatsApp corto (máximo 3-4 líneas) para que Jorge (dueño) le envíe al cliente. El mensaje debe:
- Empezar con el nombre del cliente
- Reconocer específicamente el problema que reportó (cita un detalle concreto de su feedback)
- Pedir disculpas sinceras sin sonar corporativo
- Ofrecer escuchar más en una breve llamada o mensaje (NO ofrecer descuento o compensación de entrada — eso lo decide Jorge caso a caso)
- Tono: cercano, humano, chileno informal pero respetuoso. Tutear.
- NO usar emojis excesivos. Máximo 1-2.

NO inventes datos del cliente. Si falta información, omite la cita específica y mantén general."""


def build_user_prompt(encuesta):
    """Arma el prompt con los datos del caso."""
    cliente_nombre = (
        encuesta.cliente.nombre if encuesta.cliente else encuesta.contacto_nombre
    ) or 'Cliente'

    fecha_visita = encuesta.fecha_visita.strftime('%d-%m-%Y') if encuesta.fecha_visita else 'fecha desconocida'

    parts = [f'Cliente: {cliente_nombre}']
    parts.append(f'Fecha visita: {fecha_visita}')

    if encuesta.nps_score is not None:
        parts.append(f'NPS: {encuesta.nps_score}/10 (detractor)' if encuesta.nps_score <= 6 else f'NPS: {encuesta.nps_score}/10')

    califs = []
    for label, val in [
        ('Temperatura tina', encuesta.cal_temperatura_tina),
        ('Limpieza tinas', encuesta.cal_limpieza_tinas),
        ('Limpieza cabaña', encuesta.cal_limpieza_cabana),
        ('Servicio masajes', encuesta.cal_servicio_masajes),
        ('Atención visita', encuesta.cal_atencion_visita),
        ('Experiencia general', encuesta.cal_experiencia_general),
        ('Calidad-precio', encuesta.cal_calidad_precio),
    ]:
        if val is not None and val <= 3:
            califs.append(f'{label}: {val}/5')
    if califs:
        parts.append('Calificaciones bajas: ' + '; '.join(califs))

    if encuesta.decepcion:
        parts.append(f'Lo que lo decepcionó (sus palabras): "{encuesta.decepcion}"')
    if encuesta.sugerencias:
        parts.append(f'Sugerencias que dejó: "{encuesta.sugerencias}"')
    if encuesta.lo_que_mas_gusto:
        parts.append(f'Lo que sí le gustó: "{encuesta.lo_que_mas_gusto}"')

    return '\n'.join(parts)


def generar_sugerencia_ia(encuesta):
    """Llama a Claude Sonnet via OpenRouter para sugerir mensaje WhatsApp."""
    from django.conf import settings

    try:
        from openai import OpenAI
    except ImportError:
        return '(openai SDK no instalado)'

    api_key = getattr(settings, 'OPENROUTER_API_KEY', '')
    if not api_key:
        return '(OPENROUTER_API_KEY no configurada)'

    client = OpenAI(
        api_key=api_key,
        base_url=getattr(settings, 'OPENROUTER_BASE_URL', 'https://openrouter.ai/api/v1'),
    )

    model = getattr(settings, 'REVIEW_RESPONSE_LLM_MODEL', 'anthropic/claude-sonnet-4.6')

    user_prompt = build_user_prompt(encuesta)

    try:
        response = client.chat.completions.create(
            model=model,
            messages=[
                {'role': 'system', 'content': SYSTEM_PROMPT},
                {'role': 'user', 'content': user_prompt},
            ],
            temperature=0.5,
            max_tokens=300,
        )
        return (response.choices[0].message.content or '').strip()
    except Exception as exc:
        logger.warning(f'Error generando sugerencia IA para encuesta #{encuesta.id}: {exc}')
        return f'(Error IA: {exc})'


def render_markdown(encuestas, sugerencias):
    """Construye el reporte markdown."""
    fecha = timezone.localdate().isoformat()
    lines = [
        f'# Follow-ups Pendientes — Reporte {fecha}',
        '',
        f'Total casos pendientes: **{len(encuestas)}**',
        '',
        'Cada caso incluye datos de contacto, lo que reportó el cliente y un mensaje '
        'WhatsApp sugerido por IA. Revisa, ajusta el mensaje si quieres, copia y envía '
        'directamente desde WhatsApp Web. Después marca el caso como completado en el '
        'admin Django (`/admin/ventas/encuestasatisfaccion/`).',
        '',
        '---',
        '',
    ]

    for idx, encuesta in enumerate(encuestas, 1):
        cliente_nombre = (
            encuesta.cliente.nombre if encuesta.cliente else encuesta.contacto_nombre
        ) or 'Anónimo'
        telefono = (
            encuesta.cliente.telefono if encuesta.cliente else encuesta.contacto_telefono
        ) or '(sin teléfono)'
        email = (
            encuesta.cliente.email if encuesta.cliente else encuesta.contacto_email
        ) or '(sin email)'

        fecha_visita = encuesta.fecha_visita.strftime('%d-%m-%Y') if encuesta.fecha_visita else '?'
        fecha_resp = encuesta.fecha_respuesta.strftime('%d-%m-%Y') if encuesta.fecha_respuesta else '?'

        lines.append(f'## {idx}. {cliente_nombre}')
        lines.append('')
        lines.append(f'- Encuesta ID: `{encuesta.id}`')
        lines.append(f'- Teléfono: `{telefono}`')
        lines.append(f'- Email: `{email}`')
        lines.append(f'- Fecha visita: {fecha_visita}')
        lines.append(f'- Fecha respuesta encuesta: {fecha_resp}')
        if encuesta.nps_score is not None:
            etiqueta = 'detractor' if encuesta.nps_score <= 6 else ('pasivo' if encuesta.nps_score <= 8 else 'promotor')
            lines.append(f'- NPS: **{encuesta.nps_score}/10** ({etiqueta})')

        califs_bajas = []
        for label, val in [
            ('Temperatura tina', encuesta.cal_temperatura_tina),
            ('Limpieza tinas', encuesta.cal_limpieza_tinas),
            ('Limpieza cabaña', encuesta.cal_limpieza_cabana),
            ('Servicio masajes', encuesta.cal_servicio_masajes),
            ('Atención visita', encuesta.cal_atencion_visita),
            ('Experiencia general', encuesta.cal_experiencia_general),
            ('Calidad-precio', encuesta.cal_calidad_precio),
        ]:
            if val is not None and val <= 3:
                califs_bajas.append(f'{label}: {val}/5')
        if califs_bajas:
            lines.append(f'- Calificaciones bajas: {"; ".join(califs_bajas)}')

        if encuesta.decepcion:
            lines.append(f'- **Lo decepcionó:** "{encuesta.decepcion}"')
        if encuesta.sugerencias:
            lines.append(f'- **Sugerencia:** "{encuesta.sugerencias}"')
        if encuesta.lo_que_mas_gusto:
            lines.append(f'- **Sí le gustó:** "{encuesta.lo_que_mas_gusto}"')

        lines.append('')
        lines.append('### Mensaje WhatsApp sugerido')
        lines.append('')
        sugerencia = sugerencias.get(encuesta.id, '(sin sugerencia)')
        lines.append('```')
        lines.append(sugerencia)
        lines.append('```')
        lines.append('')
        lines.append(f'[Marcar como completado en admin](https://aremko.cl/admin/ventas/encuestasatisfaccion/{encuesta.id}/change/)')
        lines.append('')
        lines.append('---')
        lines.append('')

    return '\n'.join(lines)


def markdown_to_docx_bytes(markdown_text: str) -> bytes:
    """Convierte markdown simple a .docx en memoria. Usa python-docx si está disponible."""
    try:
        from docx import Document
    except ImportError:
        return None

    doc = Document()
    for line in markdown_text.split('\n'):
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('- '):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif line.startswith('```'):
            continue  # marcador, ignorar
        elif line.strip() == '---':
            doc.add_paragraph('—' * 40)
        elif line.strip():
            doc.add_paragraph(line)
        else:
            doc.add_paragraph('')

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


class Command(BaseCommand):
    help = 'Genera y envia por email el reporte de follow-ups pendientes con sugerencias IA.'

    def add_arguments(self, parser):
        parser.add_argument('--to', default='aremkospa@gmail.com',
                            help='Email destinatario (default aremkospa@gmail.com)')
        parser.add_argument('--no-email', action='store_true',
                            help='No enviar email, solo imprimir y/o guardar.')
        parser.add_argument('--no-ai', action='store_true',
                            help='Saltar generacion IA de sugerencias (mas rapido, sin costo).')
        parser.add_argument('--output',
                            help='Si se especifica, guarda el .md en esa ruta.')
        parser.add_argument('--limit', type=int, default=0,
                            help='Si > 0, procesa solo los primeros N casos (debug).')

    def handle(self, *args, **options):
        qs = EncuestaSatisfaccion.objects.filter(
            requiere_followup=True,
            followup_completado=False,
        ).select_related('cliente').order_by('nps_score', '-fecha_respuesta')

        if options['limit'] > 0:
            qs = qs[:options['limit']]

        encuestas = list(qs)
        total = len(encuestas)

        if total == 0:
            self.stdout.write(self.style.SUCCESS('No hay follow-ups pendientes. Limpio.'))
            return

        self.stdout.write(f'{total} follow-ups pendientes encontrados.')

        sugerencias = {}
        if not options['no_ai']:
            self.stdout.write('Generando sugerencias IA (Claude Sonnet via OpenRouter)...')
            for i, encuesta in enumerate(encuestas, 1):
                self.stdout.write(f'  [{i}/{total}] caso #{encuesta.id}...', ending='\r')
                self.stdout.flush()
                sugerencias[encuesta.id] = generar_sugerencia_ia(encuesta)
            self.stdout.write('')
            self.stdout.write('Sugerencias IA generadas.')
        else:
            self.stdout.write('Saltando IA (--no-ai).')

        markdown = render_markdown(encuestas, sugerencias)

        if options['output']:
            Path(options['output']).write_text(markdown, encoding='utf-8')
            self.stdout.write(f'Markdown guardado en {options["output"]}')

        if options['no_email']:
            self.stdout.write('--no-email: skip envio.')
            return

        fecha_iso = timezone.localdate().isoformat()
        subject = f'Follow-ups Pendientes Aremko — {fecha_iso} ({total} casos)'
        body = (
            f'Reporte de los {total} clientes que dejaron feedback negativo y '
            'esperan ser contactados.\n\n'
            'Cada caso trae: datos del cliente, lo que reportó, mensaje WhatsApp sugerido por IA.\n\n'
            'Adjunto va en .docx (más fácil de leer en celular). Después de contactar a un cliente, '
            'marca el caso como completado en el admin Django.\n'
        )

        email_msg = EmailMessage(
            subject=subject,
            body=body,
            to=[options['to']],
        )

        docx_bytes = markdown_to_docx_bytes(markdown)
        if docx_bytes:
            filename = f'Aremko_Followups_{fecha_iso}.docx'
            email_msg.attach(
                filename, docx_bytes,
                'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            )
            self.stdout.write(f'Adjuntando {filename} ({len(docx_bytes)} bytes)')
        else:
            email_msg.attach(f'Aremko_Followups_{fecha_iso}.md', markdown.encode('utf-8'), 'text/markdown')
            self.stdout.write('python-docx no disponible, adjuntando .md')

        try:
            email_msg.send(fail_silently=False)
            self.stdout.write(self.style.SUCCESS(f'Email enviado a {options["to"]}'))
        except Exception as exc:
            self.stderr.write(self.style.ERROR(f'Error enviando email: {exc}'))
            raise
