"""Analisis IA de snapshots Meta (FB + IG + Ads).

Toma un snapshot generado por meta_reporter.py y genera un analisis
ejecutivo con Claude Sonnet via OpenRouter. Output estructurado JSON
con resumen, alertas, oportunidades y acciones recomendadas.

Usado por:
- Boton admin "Snapshot + Analisis + Email" para diagnostico ad-hoc
- Brief semanal automatico (lunes 10am) para enriquecer el contexto
"""
import json
import logging
from datetime import date
from typing import Optional

from django.conf import settings

logger = logging.getLogger(__name__)


SYSTEM_PROMPT = """Eres analista de marketing digital especializado en redes sociales y paid media para PyMEs chilenas del sector turismo/wellness.

Tu cliente es Aremko Spa Boutique en Puerto Varas (Chile). Negocio:
- 53k fans en Facebook, 59k followers en Instagram
- Vende: cabañas + tinas calientes + masajes (pack 3-en-1, $190-220k CLP)
- Diferencial: sonido del río Pescado, 48 paneles solares, sin leña
- Buyer persona principal: pareja en burnout 30-45 años (70% del público)
- Frase ancla: "2 días aquí valen una semana de vacaciones"

Vas a recibir un snapshot de 28 días con 3 fuentes:
1. **Facebook orgánico**: posts publicados, engagement por post
2. **Instagram orgánico**: alcance, follower growth, top media, interactions
3. **Paid Ads**: TODAS las cuentas publicitarias accesibles (puede haber varias). Cada una con summary, insights del periodo, y campañas con su rendimiento.

Tu tarea: genera un análisis JSON ejecutivo con esta estructura EXACTA:

{
  "resumen_ejecutivo": "3-4 frases. Foco en qué está pasando ahora vs cómo debería estar. Concreto, sin adjetivos vacíos. Mencionar números reales del snapshot.",

  "metricas_destacadas": {
    "facebook": {
      "fans": <número>,
      "posts_28d": <número>,
      "engagement_total_28d": <número>,
      "engagement_por_post_promedio": <número>,
      "diagnostico_breve": "1 frase con diagnóstico (ej: 'engagement bajo para 53k fans, frecuencia de 1 post cada 9 días')"
    },
    "instagram": {
      "followers": <número>,
      "alcance_28d": <número>,
      "interacciones_28d": <número>,
      "ratio_alcance_followers": "<número con 1 decimal>",
      "diagnostico_breve": "1 frase"
    },
    "paid_ads": {
      "cuentas_activas_28d": <número de cuentas con spend > 0>,
      "spend_total_28d_clp": <suma en CLP>,
      "campañas_activas": <número total de campañas con effective_status=ACTIVE>,
      "diagnostico_breve": "1 frase"
    }
  },

  "alertas_criticas": [
    "Lista de 1-3 alertas urgentes detectadas. Ej: 'Cero conversaciones de Messenger en 28 días con $50k de inversión = $50k por lead, no rentable'. Solo cosas urgentes, no decoración."
  ],

  "oportunidades_detectadas": [
    "Lista de 2-4 oportunidades concretas. Ej: 'Reels orgánicos del 27-abr (4 reactions FB, 10 likes IG) son los más performantes — replicar formato y boostearlos esta semana'."
  ],

  "comparacion_organico_vs_paid": "2-3 frases sobre la relación entre el contenido orgánico y los boosts. Ej: '95% del engagement viene de paid (12k vs 666 orgánico) — orgánico está fallando como canal independiente, paid lo está supliendo pero a costo X por interacción'.",

  "acciones_recomendadas_esta_semana": [
    {
      "accion": "Acción concreta 1",
      "responsable_sugerido": "Jorge | Daniela | Equipo",
      "esfuerzo": "bajo | medio | alto",
      "impacto_esperado": "1 frase",
      "metrica_a_revisar_proxima_semana": "Métrica específica para validar si funcionó"
    }
  ],

  "preguntas_para_revisar": [
    "Lista de 2-3 preguntas que Jorge debería responder. Ej: '¿Las campañas con objective=MESSAGES están convirtiendo a reservas reales o solo a mensajes que no avanzan? Sin GA4 conversion tracking conectado, no podemos saberlo.'"
  ]
}

REGLAS DE GENERACIÓN:
- Español latinoamericano (NUNCA voseo argentino: usa "tú", "puedes", "quieres", no "vos podés")
- Cero adjetivos vacíos ("excelente", "increíble", "magia"). Solo datos.
- Si no hay datos suficientes para un campo, escribir "(sin datos)" no inventar
- Las acciones recomendadas deben ser EJECUTABLES esta semana, no proyectos a 3 meses
- JSON estricto, sin markdown wrapper, sin comentarios"""


def _strip_fences(raw: str) -> str:
    s = (raw or '').strip()
    if s.startswith('```'):
        lines = s.split('\n')
        if lines and lines[0].startswith('```'):
            lines = lines[1:]
        if lines and lines[-1].strip() == '```':
            lines = lines[:-1]
        return '\n'.join(lines).strip()
    return s


def analyze_snapshot(snapshot: dict, model: Optional[str] = None) -> dict:
    """Toma un snapshot Meta y devuelve un analisis IA estructurado.

    Args:
        snapshot: dict generado por meta_reporter.get_full_snapshot()
        model: override del modelo LLM (default: claude-sonnet-4.6 via OpenRouter)

    Returns:
        dict con keys: resumen_ejecutivo, metricas_destacadas, alertas_criticas,
        oportunidades_detectadas, comparacion_organico_vs_paid,
        acciones_recomendadas_esta_semana, preguntas_para_revisar.

    Raises:
        ValueError si OPENROUTER_API_KEY no esta configurada o si el LLM
        no devuelve JSON valido.
    """
    from openai import OpenAI

    api_key = getattr(settings, 'OPENROUTER_API_KEY', '')
    base_url = getattr(settings, 'OPENROUTER_BASE_URL', 'https://openrouter.ai/api/v1')
    if not api_key:
        raise ValueError('OPENROUTER_API_KEY no configurada')

    model = model or getattr(settings, 'META_ANALYSIS_LLM_MODEL', 'anthropic/claude-sonnet-4.6')

    fecha_actual = date.today().isoformat()

    user_prompt = (
        f"Fecha actual: {fecha_actual}\n\n"
        f"Snapshot de los ultimos {snapshot.get('period_days', 28)} dias:\n\n"
        f"{json.dumps(snapshot, indent=2, ensure_ascii=False, default=str)[:18000]}"
    )

    client = OpenAI(api_key=api_key, base_url=base_url)
    logger.info(f'Llamando a {model} para analisis IA del snapshot Meta')

    response = client.chat.completions.create(
        model=model,
        messages=[
            {'role': 'system', 'content': SYSTEM_PROMPT},
            {'role': 'user', 'content': user_prompt},
        ],
        temperature=0.3,
        max_tokens=4000,
        response_format={'type': 'json_object'},
    )

    raw = response.choices[0].message.content or ''
    cleaned = _strip_fences(raw)
    try:
        analysis = json.loads(cleaned)
    except json.JSONDecodeError as exc:
        logger.error(f'LLM analisis no devolvio JSON valido. Raw[:300]: {raw[:300]}')
        raise ValueError(f'LLM analisis no es JSON valido: {exc}')

    return analysis


def render_analysis_to_markdown(analysis: dict, snapshot: dict) -> str:
    """Convierte el analisis JSON a markdown legible para email/docx."""
    period_days = snapshot.get('period_days', 28)
    fecha = date.today().isoformat()

    lines = [
        f'# Diagnóstico Meta Aremko — {fecha}',
        '',
        f'Análisis de los últimos **{period_days} días** de Facebook + Instagram + Ads (todas las cuentas).',
        '',
        '## Resumen ejecutivo',
        '',
        analysis.get('resumen_ejecutivo', '(sin resumen)'),
        '',
        '## Métricas destacadas',
        '',
    ]

    metricas = analysis.get('metricas_destacadas', {})

    fb = metricas.get('facebook', {})
    if fb:
        lines.extend([
            '### Facebook',
            f"- Fans: **{fb.get('fans', 0):,}**",
            f"- Posts en {period_days}d: **{fb.get('posts_28d', 0)}**",
            f"- Engagement total: **{fb.get('engagement_total_28d', 0)}**",
            f"- Engagement por post promedio: **{fb.get('engagement_por_post_promedio', 0)}**",
            f"- Diagnóstico: {fb.get('diagnostico_breve', '')}",
            '',
        ])

    ig = metricas.get('instagram', {})
    if ig:
        lines.extend([
            '### Instagram',
            f"- Followers: **{ig.get('followers', 0):,}**",
            f"- Alcance {period_days}d: **{ig.get('alcance_28d', 0):,}**",
            f"- Interacciones {period_days}d: **{ig.get('interacciones_28d', 0):,}**",
            f"- Ratio alcance/followers: **{ig.get('ratio_alcance_followers', 'n/d')}x**",
            f"- Diagnóstico: {ig.get('diagnostico_breve', '')}",
            '',
        ])

    paid = metricas.get('paid_ads', {})
    if paid:
        lines.extend([
            '### Paid Ads',
            f"- Cuentas con actividad: **{paid.get('cuentas_activas_28d', 0)}**",
            f"- Inversión total: **${paid.get('spend_total_28d_clp', 0):,} CLP**",
            f"- Campañas activas: **{paid.get('campañas_activas', 0)}**",
            f"- Diagnóstico: {paid.get('diagnostico_breve', '')}",
            '',
        ])

    alertas = analysis.get('alertas_criticas', [])
    if alertas:
        lines.append('## Alertas críticas')
        lines.append('')
        for a in alertas:
            lines.append(f'- ⚠️ {a}')
        lines.append('')

    oportunidades = analysis.get('oportunidades_detectadas', [])
    if oportunidades:
        lines.append('## Oportunidades detectadas')
        lines.append('')
        for o in oportunidades:
            lines.append(f'- 💡 {o}')
        lines.append('')

    comparacion = analysis.get('comparacion_organico_vs_paid', '')
    if comparacion:
        lines.extend([
            '## Orgánico vs Paid',
            '',
            comparacion,
            '',
        ])

    acciones = analysis.get('acciones_recomendadas_esta_semana', [])
    if acciones:
        lines.append('## Acciones recomendadas esta semana')
        lines.append('')
        for i, a in enumerate(acciones, 1):
            lines.append(f"### {i}. {a.get('accion', '')}")
            lines.append(f"- **Responsable**: {a.get('responsable_sugerido', '?')}")
            lines.append(f"- **Esfuerzo**: {a.get('esfuerzo', '?')}")
            lines.append(f"- **Impacto esperado**: {a.get('impacto_esperado', '')}")
            lines.append(f"- **Métrica a revisar**: {a.get('metrica_a_revisar_proxima_semana', '')}")
            lines.append('')

    preguntas = analysis.get('preguntas_para_revisar', [])
    if preguntas:
        lines.append('## Preguntas para revisar')
        lines.append('')
        for p in preguntas:
            lines.append(f'- ❓ {p}')
        lines.append('')

    return '\n'.join(lines)


def markdown_to_docx_bytes(md: str) -> Optional[bytes]:
    """Convierte markdown del analisis a .docx en memoria. None si falla."""
    try:
        from docx import Document
    except ImportError:
        return None

    from io import BytesIO
    doc = Document()
    for line in md.split('\n'):
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('- '):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif line.strip():
            doc.add_paragraph(line)
        else:
            doc.add_paragraph('')
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def send_analysis_by_email(analysis: dict, snapshot: dict, to_emails: Optional[list] = None) -> bool:
    """Envia el analisis IA por email (default: 3 emails internos Aremko)."""
    from django.core.mail import EmailMessage

    if not to_emails:
        to_emails = [
            'aremkospa@gmail.com',
            'abonosaremko@gmail.com',
            'atoloza1970@gmail.com',
        ]

    fecha = date.today().isoformat()
    md = render_analysis_to_markdown(analysis, snapshot)

    subject = f'📊 Diagnóstico Meta Aremko — {fecha}'
    body = (
        f'Diagnóstico automático de redes sociales y paid ads de Aremko al {fecha}.\n\n'
        f'Generado bajo demanda desde el admin Django. Datos de los últimos '
        f'{snapshot.get("period_days", 28)} días.\n\n'
        f'Adjunto va el reporte completo en .docx con resumen ejecutivo, '
        f'métricas, alertas, oportunidades y acciones recomendadas.\n'
    )

    msg = EmailMessage(subject=subject, body=body, to=to_emails)

    docx_bytes = markdown_to_docx_bytes(md)
    if docx_bytes:
        filename = f'Aremko_Diagnostico_Meta_{fecha}.docx'
        msg.attach(
            filename, docx_bytes,
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        )
    else:
        msg.attach(f'Aremko_Diagnostico_Meta_{fecha}.md', md.encode('utf-8'), 'text/markdown')

    try:
        msg.send(fail_silently=False)
        return True
    except Exception as exc:
        logger.error(f'Error enviando diagnostico Meta: {exc}')
        return False
