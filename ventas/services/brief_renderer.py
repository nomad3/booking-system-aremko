"""Renderiza el output JSON del brief semanal a markdown extenso y .docx.

El brief semanal es un documento operativo grande. Se manda por email con
3 partes:
- Resumen ejecutivo en el cuerpo HTML del email (lectura rápida en celular)
- Markdown completo como vista previa en email body
- .docx adjunto con todo el contenido formateado para imprimir/distribuir
"""
import logging
from datetime import date
from io import BytesIO
from typing import Optional

logger = logging.getLogger(__name__)


def _bullet(items, prefix: str = '- ') -> list:
    if not items:
        return []
    return [f'{prefix}{item}' for item in items]


def render_brief_to_markdown(brief: dict, semana_inicio: date, semana_fin: date) -> str:
    """Convierte el output JSON del brief a markdown extenso y legible."""
    lines = []
    lines.append(f'# Brief Marketing Aremko — Semana del {semana_inicio.strftime("%d-%m-%Y")} al {semana_fin.strftime("%d-%m-%Y")}')
    lines.append('')

    # === Resumen ejecutivo ===
    res = brief.get('resumen_ejecutivo', {})
    if res:
        lines.append('## 📋 Resumen ejecutivo')
        lines.append('')
        if res.get('headline'):
            lines.append(f'**{res["headline"]}**')
            lines.append('')
        if res.get('alerta_critica_si_existe'):
            lines.append(f'> ⚠️ **Alerta crítica**: {res["alerta_critica_si_existe"]}')
            lines.append('')
        if res.get('objetivo_de_la_semana'):
            lines.append(f'**Objetivo de la semana**: {res["objetivo_de_la_semana"]}')
            lines.append('')
        if res.get('puntos_clave'):
            lines.append('**Puntos clave**:')
            lines.extend(_bullet(res['puntos_clave']))
            lines.append('')

    # === Fechas clave próximas ===
    fechas = brief.get('fechas_clave_proximas_4_semanas', [])
    if fechas:
        lines.append('## 📅 Fechas clave próximas 4 semanas')
        lines.append('')
        for f in fechas:
            lines.append(f'- **{f.get("fecha", "?")}** — {f.get("evento", "?")}: {f.get("implicancia_marketing", "")}')
        lines.append('')

    # === Diagnóstico extenso ===
    diag = brief.get('diagnostico_extenso', {})
    if diag:
        lines.append('## 🔍 Diagnóstico extenso')
        lines.append('')

        # Voz del cliente
        voc = diag.get('voz_del_cliente', {})
        if voc:
            lines.append('### 🗣️ Voz del cliente')
            lines.append('')
            if voc.get('sentimiento_general'):
                lines.append(voc['sentimiento_general'])
                lines.append('')
            if voc.get('alertas_operativas'):
                lines.append('**Alertas operativas detectadas**:')
                for a in voc['alertas_operativas']:
                    lines.append(f'- ⚠️ {a.get("alerta", "")} — *fuente: {a.get("fuente", "?")}, frecuencia: {a.get("frecuencia", "?")}* — Acción: {a.get("accion_sugerida", "")}')
                lines.append('')
            if voc.get('oportunidades_de_contenido'):
                lines.append('**Oportunidades de contenido (para Reels/posts)**:')
                lines.extend(_bullet(voc['oportunidades_de_contenido']))
                lines.append('')
            if voc.get('follow_ups_urgentes_count'):
                lines.append(f'**Follow-ups urgentes pendientes**: {voc["follow_ups_urgentes_count"]}')
                lines.append('')

        # Redes sociales
        redes = diag.get('redes_sociales', {})
        if redes:
            lines.append('### 📱 Redes sociales')
            lines.append('')
            if redes.get('facebook_diagnostico'):
                lines.append(f'**Facebook**: {redes["facebook_diagnostico"]}')
                lines.append('')
            if redes.get('instagram_diagnostico'):
                lines.append(f'**Instagram**: {redes["instagram_diagnostico"]}')
                lines.append('')
            if redes.get('paid_ads_diagnostico'):
                lines.append(f'**Paid Ads**: {redes["paid_ads_diagnostico"]}')
                lines.append('')
            if redes.get('comparacion_organico_vs_paid'):
                lines.append(f'**Orgánico vs Paid**: {redes["comparacion_organico_vs_paid"]}')
                lines.append('')
            if redes.get('top_3_aprendizajes_de_contenido'):
                lines.append('**Top 3 aprendizajes de contenido**:')
                lines.extend(_bullet(redes['top_3_aprendizajes_de_contenido']))
                lines.append('')

        # Web y SEO
        web = diag.get('web_y_seo', {})
        if web:
            lines.append('### 🌐 Web y SEO')
            lines.append('')
            if web.get('trafico_resumen'):
                lines.append(f'**Tráfico**: {web["trafico_resumen"]}')
                lines.append('')
            if web.get('embudo_reservas'):
                lines.append(f'**Embudo reservas**: {web["embudo_reservas"]}')
                lines.append('')
            if web.get('seo_resumen'):
                lines.append(f'**SEO**: {web["seo_resumen"]}')
                lines.append('')
            if web.get('top_3_acciones_seo_esta_semana'):
                lines.append('**Top 3 acciones SEO esta semana**:')
                lines.extend(_bullet(web['top_3_acciones_seo_esta_semana']))
                lines.append('')

        # Comercial y pipeline
        com = diag.get('comercial_y_pipeline', {})
        if com:
            lines.append('### 💰 Comercial y pipeline')
            lines.append('')
            # Nuevos campos extendidos
            for label, key in [
                ('Reservas última semana', 'reservas_resumen'),
                ('Tendencia vs semana anterior', 'tendencia_vs_semana_anterior'),
                ('Resumen mensual', 'comparacion_mes'),
                ('Mezcla por categoría', 'mezcla_por_categoria'),
                ('Packs vs reservas simples', 'analisis_packs_vs_simples'),
                ('Canales de pago y cambios', 'canales_pago_y_cambios'),
                ('Comportamiento cliente', 'comportamiento_cliente'),
                ('Disponibilidad próxima semana', 'disponibilidad_proxima_semana'),
                ('Abandonos Flow', 'abandonos_flow'),
                ('Implicancia comunicacional', 'implicancia_comunicacional'),
            ]:
                if com.get(key):
                    lines.append(f'**{label}**: {com[key]}')
                    lines.append('')
            # Compat con formato antiguo (mantener por si)
            if com.get('tasa_conversion_pago') and not com.get('tendencia_vs_semana_anterior'):
                lines.append(f'**Tasa conversión a pago**: {com["tasa_conversion_pago"]}')
                lines.append('')
            if com.get('servicios_top') and not com.get('mezcla_por_categoria'):
                lines.append(f'**Top servicios**: {com["servicios_top"]}')
                lines.append('')

    # === Calendario semanal ===
    cal = brief.get('calendario_semanal', [])
    if cal:
        lines.append('## 📆 Calendario semanal')
        lines.append('')
        for d in cal:
            pubs = d.get('publicaciones', [])
            if pubs:
                lines.append(f'### {d.get("dia", "")} {d.get("fecha", "")}')
                for p in pubs:
                    lines.append(f'- **{p.get("hora", "?")}** [{p.get("canal", "")}] {p.get("tipo", "")}: {p.get("concepto_corto", "")} — *{p.get("estado", "")}*')
                lines.append('')

    # === Drafts completos ===
    drafts = brief.get('drafts_completos', {})
    if drafts:
        lines.append('## ✍️ Drafts completos (listos para copiar/pegar)')
        lines.append('')

        # GBP post
        gbp = drafts.get('gbp_post', {})
        if gbp.get('necesario_esta_semana'):
            lines.append('### Google Business Profile')
            lines.append(f'**Responsable**: {gbp.get("responsable", "?")} · **Tiempo**: {gbp.get("tiempo_estimado", "?")}')
            lines.append('')
            lines.append('```')
            lines.append(gbp.get('texto', ''))
            lines.append('```')
            lines.append('')
            if gbp.get('url_cta'):
                lines.append(f'**URL CTA**: {gbp["url_cta"]}')
            if gbp.get('foto_sugerida'):
                lines.append(f'**Foto sugerida**: {gbp["foto_sugerida"]}')
            lines.append('')

        # Reel martes
        reel_m = drafts.get('reel_martes', {})
        if reel_m.get('necesario_esta_semana'):
            _render_reel(lines, reel_m, 'Reel martes')

        # Carrusel miércoles
        car = drafts.get('carrusel_miercoles', {})
        if car.get('necesario_esta_semana'):
            lines.append('### Carrusel miércoles (Instagram)')
            lines.append(f'**Responsable**: {car.get("responsable", "?")} · **Tiempo**: {car.get("tiempo_estimado", "?")}')
            lines.append('')
            lines.append(f'**Concepto**: {car.get("concepto", "")}')
            lines.append('')
            for s in car.get('slides', []):
                lines.append(f'- **Slide {s.get("numero", "?")}** ({s.get("rol", "")}): {s.get("imagen_sugerida", "")} | overlay: "{s.get("texto_overlay", "")}"')
            lines.append('')
            if car.get('caption_completo'):
                lines.append('**Caption completo**:')
                lines.append('```')
                lines.append(car['caption_completo'])
                lines.append('```')
                lines.append('')

        # Reel jueves
        reel_j = drafts.get('reel_jueves', {})
        if reel_j.get('necesario_esta_semana'):
            _render_reel(lines, reel_j, 'Reel jueves')

        # Stories diarias
        stories = drafts.get('stories_diarias', [])
        if stories:
            lines.append('### Stories diarias')
            lines.append('')
            for s in stories:
                lines.append(f'- **{s.get("dia", "?")}** [{s.get("tipo", "")}]: {s.get("concepto", "")} — *texto: {s.get("texto_sugerido", "")}*')
            lines.append('')

        # Email engaged
        email_e = drafts.get('email_engaged', {})
        if email_e.get('necesario_esta_semana'):
            lines.append('### Email a segmento engaged')
            lines.append(f'**Responsable**: {email_e.get("responsable", "?")} · **Tiempo**: {email_e.get("tiempo_estimado", "?")} · **Segmento**: {email_e.get("segmento_destinatario", "?")}')
            lines.append('')
            if email_e.get('asunto'):
                lines.append(f'**Asunto**: {email_e["asunto"]}')
            if email_e.get('preheader'):
                lines.append(f'**Preheader**: {email_e["preheader"]}')
            lines.append('')
            if email_e.get('cuerpo_html_resumen'):
                lines.append(f'**Estructura HTML**: {email_e["cuerpo_html_resumen"]}')
                lines.append('')
            if email_e.get('cuerpo_texto_plano_completo'):
                lines.append('**Cuerpo texto plano**:')
                lines.append('```')
                lines.append(email_e['cuerpo_texto_plano_completo'])
                lines.append('```')
                lines.append('')

        # Post blog si aplica
        blog = drafts.get('post_blog_si_aplica', {})
        if blog.get('necesario_esta_semana'):
            lines.append('### Post de blog esta semana')
            lines.append(f'**Tema**: {blog.get("tema_sugerido", "?")} · **Responsable**: {blog.get("responsable", "?")}')
            lines.append('')

    # === Ideas próximas 2 semanas ===
    ideas = brief.get('ideas_contenido_proximas_2_semanas', [])
    if ideas:
        lines.append('## 💡 Ideas de contenido para próximas 2 semanas')
        lines.append('')
        for i in ideas:
            lines.append(f'- **{i.get("semana", "?")}** [{i.get("tipo", "")}]: {i.get("tema", "")} — *{i.get("razon", "")}*')
        lines.append('')

    # === Acciones paid ads ===
    paid_actions = brief.get('acciones_paid_ads_recomendadas', [])
    if paid_actions:
        lines.append('## 💰 Acciones paid ads recomendadas')
        lines.append('')
        for a in paid_actions:
            lines.append(f'### {a.get("accion", "")}')
            lines.append(f'- **Responsable**: {a.get("responsable", "?")} · **Esfuerzo**: {a.get("esfuerzo", "?")}')
            lines.append(f'- **Razón**: {a.get("razon", "")}')
            lines.append(f'- **Métrica a mover**: {a.get("metrica_a_mover", "")}')
            lines.append('')

    # === Alertas operativas no marketing ===
    alertas_op = brief.get('alertas_operativas_no_marketing', [])
    if alertas_op:
        lines.append('## 🚨 Alertas operativas (no marketing)')
        lines.append('')
        lines.extend(_bullet(alertas_op))
        lines.append('')

    # === Métricas viernes ===
    metricas = brief.get('metricas_a_revisar_viernes', [])
    if metricas:
        lines.append('## 📊 Métricas a revisar viernes')
        lines.append('')
        lines.extend(_bullet(metricas))
        lines.append('')

    # === Preguntas para Jorge ===
    preg = brief.get('preguntas_abiertas_para_jorge', [])
    if preg:
        lines.append('## ❓ Preguntas para Jorge')
        lines.append('')
        for p in preg:
            lines.append(f'- {p}')
        lines.append('')

    # === Recordatorios ===
    rec = brief.get('recordatorios', [])
    if rec:
        lines.append('## 🔔 Recordatorios')
        lines.append('')
        lines.extend(_bullet(rec))
        lines.append('')

    return '\n'.join(lines)


def _render_reel(lines: list, reel: dict, titulo: str):
    """Helper para renderizar un Reel."""
    lines.append(f'### {titulo} (Instagram)')
    lines.append(f'**Responsable**: {reel.get("responsable", "?")} · **Tiempo**: {reel.get("tiempo_estimado", "?")} · **Duración**: {reel.get("duracion_objetivo_segundos", "?")}s')
    lines.append('')
    if reel.get('concepto'):
        lines.append(f'**Concepto**: {reel["concepto"]}')
        lines.append('')
    if reel.get('filtro_5_50'):
        lines.append(f'**Filtro 5/50 (Víctor Eras)**: {reel["filtro_5_50"]}')
        lines.append('')

    # Guion
    guion = reel.get('guion', [])
    if guion:
        lines.append('**Guión**:')
        for b in guion:
            lines.append(f'- **{b.get("bloque", "?")}**: {b.get("texto", "")}')
        lines.append('')

    if reel.get('tomas_sugeridas'):
        lines.append('**Tomas sugeridas**:')
        lines.extend(_bullet(reel['tomas_sugeridas']))
        lines.append('')

    if reel.get('audio_sugerido'):
        lines.append(f'**Audio**: {reel["audio_sugerido"]}')
        lines.append('')

    if reel.get('caption_completo'):
        lines.append('**Caption completo**:')
        lines.append('```')
        lines.append(reel['caption_completo'])
        lines.append('```')
        lines.append('')

    if reel.get('hashtags'):
        lines.append(f'**Hashtags**: {" ".join(reel["hashtags"])}')
        lines.append('')


def render_brief_to_docx(brief: dict, semana_inicio: date, semana_fin: date) -> Optional[bytes]:
    """Convierte el brief a .docx en memoria. None si falla."""
    try:
        from docx import Document
    except ImportError:
        return None

    md = render_brief_to_markdown(brief, semana_inicio, semana_fin)
    doc = Document()

    in_code_block = False
    for line in md.split('\n'):
        if line.strip() == '```':
            in_code_block = not in_code_block
            continue
        if in_code_block:
            p = doc.add_paragraph(line)
            for run in p.runs:
                run.font.name = 'Courier New'
            continue
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('- '):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif line.startswith('> '):
            p = doc.add_paragraph(line[2:])
            p.paragraph_format.left_indent = None
        elif line.strip():
            doc.add_paragraph(line)
        else:
            doc.add_paragraph('')

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def render_executive_summary_html(brief: dict) -> str:
    """Renderiza solo el resumen ejecutivo + objetivo en HTML para el body
    del email. El resto del brief va en el .docx adjunto.
    """
    res = brief.get('resumen_ejecutivo', {}) or {}
    headline = res.get('headline', '')
    objetivo = res.get('objetivo_de_la_semana', '')
    puntos = res.get('puntos_clave', []) or []
    alerta = res.get('alerta_critica_si_existe', '')

    parts = ['<div style="font-family:Arial,sans-serif;color:#222;">']
    if headline:
        parts.append(f'<h2 style="color:#1a73e8;">{headline}</h2>')
    if alerta:
        parts.append(
            f'<div style="background:#fff3cd;border-left:4px solid #ff9800;'
            f'padding:12px 16px;margin:16px 0;border-radius:4px;">'
            f'<strong>⚠️ Alerta crítica</strong><br>{alerta}</div>'
        )
    if objetivo:
        parts.append(
            f'<div style="background:#e8f5e9;border-left:4px solid #2e7d32;'
            f'padding:12px 16px;margin:16px 0;border-radius:4px;">'
            f'<strong>🎯 Objetivo de la semana</strong><br>{objetivo}</div>'
        )
    if puntos:
        parts.append('<h3>Puntos clave</h3><ul>')
        for p in puntos:
            parts.append(f'<li style="margin-bottom:8px;">{p}</li>')
        parts.append('</ul>')

    parts.append(
        '<hr style="margin:24px 0;border:none;border-top:1px solid #ddd;">'
        '<p style="color:#666;font-size:14px;">'
        'El brief completo (diagnóstico extenso + drafts listos + acciones recomendadas) '
        'va en el .docx adjunto.</p>'
    )
    parts.append('</div>')
    return ''.join(parts)
